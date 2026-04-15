import os
import uuid
import shutil
import subprocess
import json
import logging
import re
from collections import Counter
from datetime import datetime, timedelta
from tempfile import TemporaryDirectory
from typing import List
from urllib.parse import quote_plus
import uvicorn
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import JSONResponse, HTMLResponse
from pymongo import MongoClient
from pydub import AudioSegment
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import time
import asyncio
import torch
import boto3
from botocore.exceptions import NoCredentialsError
from openai import OpenAI
from groq import Groq


client = OpenAI()
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# === GPU CHECK ===
print("Using GPU:", torch.cuda.is_available())

# === INIT ===
app = FastAPI()
openai.api_key = os.getenv("OPENAI_API_KEY")

# === LOGGING ===
logger = logging.getLogger("video_processor")
logging.basicConfig(level=logging.INFO)
logging.getLogger("httpx").setLevel(logging.WARNING)

# === AWS CONFIG ===
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "ap-south-1")
AWS_S3_BUCKET = "imeetpro-225220763325"

s3_client = boto3.client(
    "s3",
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
    region_name=AWS_REGION
)

# === MONGO DB ===
mongo_user = quote_plus("connectly")
mongo_password = quote_plus("LT@connect25")
mongo_host = "192.168.48.201"
mongo_port = "27017"

MONGO_URI = (
    f"mongodb://{mongo_user}:{mongo_password}@"
    f"{mongo_host}:{mongo_port}/test2?authSource=admin"
)

mongo_client = MongoClient(MONGO_URI)
db = mongo_client["test2"]
dev_collection     = db["Developer"]
non_dev_collection = db["Non-Developer"]


# ══════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════

def clean_markdown(text: str) -> str:
    text = re.sub(r"```(\w+)?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"^### (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^## (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^# (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()

def upload_to_s3(folder: str, file_path: str, file_name: str):
    try:
        s3_key = f"{folder}/{file_name}"
        s3_client.upload_file(file_path, AWS_S3_BUCKET, s3_key)
        url = f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        logger.info(f"[S3] Uploaded {file_name} -> {url}")
        return url
    except Exception as e:
        logger.error(f"S3 upload failed: {e}")
        raise

def format_srt_time(seconds: float) -> str:
    td = timedelta(seconds=seconds)
    total = int(td.total_seconds())
    millis = int((td.total_seconds() - total) * 1000)
    return f"{str(timedelta(seconds=total)).zfill(8)},{millis:03}"

def create_srt_from_segments(segments: List[dict], output_path: str):
    if not segments:
        logger.warning(f"[WARN] No segments to write for {output_path}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("1\n00:00:00,000 --> 00:00:05,000\n[No speech detected]\n\n")
        return
    logger.info(f"[INFO] Writing subtitles to {output_path} ({len(segments)} segments)")
    with open(output_path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, start=1):
            start = format_srt_time(seg.get("start", 0))
            end = format_srt_time(seg.get("end", 0))
            text = (seg.get("text") or "").strip()
            if not text:
                text = "[Silence]"
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

def generate_graph(dot_code: str, output_path: str):
    from graphviz import Source
    s = Source(dot_code)
    return s.render(filename=output_path, format="png", cleanup=True)

def generate_summary_pdf(summary_text: str, domain: str, output_path: str) -> bool:
    """
    Generates a styled PDF from the class notebook summary.
    Returns True on success, False on failure.
    Logs every step so failures are visible in server logs.
    """
    try:
        logger.info(f"[PDF] Starting generation | domain={domain} | output={output_path}")
        logger.info(f"[PDF] Summary text length: {len(summary_text)} chars")

        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        doc = SimpleDocTemplate(
            output_path, pagesize=letter,
            rightMargin=0.75*inch, leftMargin=0.75*inch,
            topMargin=0.75*inch, bottomMargin=0.75*inch
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('CTitle', parent=styles['Title'],
            fontSize=20, textColor=colors.HexColor('#1F5C99'),
            spaceAfter=4, spaceBefore=0, alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('CSubtitle', parent=styles['Normal'],
            fontSize=11, textColor=colors.HexColor('#555555'),
            spaceAfter=16, alignment=TA_CENTER)
        h1_style = ParagraphStyle('CH1', parent=styles['Normal'],
            fontSize=14, textColor=colors.HexColor('#1F5C99'),
            fontName='Helvetica-Bold', spaceBefore=16, spaceAfter=5)
        body_style = ParagraphStyle('CBody', parent=styles['Normal'],
            fontSize=10, leading=15, spaceAfter=7,
            alignment=TA_JUSTIFY, textColor=colors.HexColor('#222222'))
        bullet_style = ParagraphStyle('CBullet', parent=styles['Normal'],
            fontSize=10, leading=14, spaceAfter=4,
            leftIndent=18, textColor=colors.HexColor('#222222'))
        qa_q_style = ParagraphStyle('CQQ', parent=styles['Normal'],
            fontSize=10, fontName='Helvetica-Bold', leading=14,
            spaceAfter=3, textColor=colors.HexColor('#1F5C99'), spaceBefore=10)
        qa_a_style = ParagraphStyle('CQA', parent=styles['Normal'],
            fontSize=10, leading=14, spaceAfter=7,
            leftIndent=14, textColor=colors.HexColor('#333333'))
        subhead_style = ParagraphStyle('CSubHead', parent=styles['Normal'],
            fontSize=11, fontName='Helvetica-Bold',
            textColor=colors.HexColor('#2E75B6'),
            spaceBefore=10, spaceAfter=3)
        takeaway_style = ParagraphStyle('CTakeaway', parent=styles['Normal'],
            fontSize=10, leading=14, spaceAfter=5,
            leftIndent=10, textColor=colors.HexColor('#1B5E20'))
        footer_style = ParagraphStyle('CFooter', parent=styles['Normal'],
            fontSize=8, textColor=colors.HexColor('#888888'),
            alignment=TA_CENTER, spaceBefore=5)
        code_style = ParagraphStyle('CCode', parent=styles['Normal'],
            fontSize=9, leading=13, spaceAfter=5,
            fontName='Courier', textColor=colors.HexColor('#1a1a2e'),
            backColor=colors.HexColor('#f4f4f4'),
            leftIndent=12, rightIndent=12, borderPad=4)
        tcode_style = ParagraphStyle('CTCode', parent=styles['Normal'],
            fontSize=10, leading=14, spaceAfter=5,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#0d47a1'),
            leftIndent=10)

        def sec_hr():
            return HRFlowable(width="100%", thickness=2,
                color=colors.HexColor('#1F5C99'), spaceAfter=6, spaceBefore=2)
        def light_hr():
            return HRFlowable(width="100%", thickness=0.5,
                color=colors.HexColor('#DDDDDD'), spaceAfter=5, spaceBefore=2)
        def safe(text):
            return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        SECTION_HEADINGS = {
            "Overview", "Today's Session Notes", "Core Concepts",
            "Real-Time Examples", "Practical Reference",
            "Failure Signals", "Common Errors",
            "Common Mistakes Students Make", "Key Takeaways",
            "Interview Questions", "Terminology Glossary",
            "Escalation Guide", "Best Practices"
        }
        SUBHEAD_SECTIONS = {
            "Common Errors", "Common Mistakes Students Make",
            "Escalation Guide", "Best Practices",
            "Core Concepts", "Terminology Glossary",
            "Practical Reference"
        }

        story = []
        safe_domain = safe(domain)
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(safe_domain + " — Class Notes", title_style))
        story.append(Paragraph("iMentora Training Session", subtitle_style))
        story.append(sec_hr())
        story.append(Spacer(1, 0.1*inch))

        lines = summary_text.split('\n')
        current_section = ""

        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line in SECTION_HEADINGS:
                current_section = line
                story.append(Paragraph(safe(line), h1_style))
                story.append(sec_hr())
                continue
            if current_section == "Interview Questions":
                if line.startswith("Q") and (":" in line or "?" in line):
                    story.append(Paragraph(safe(line), qa_q_style))
                else:
                    story.append(Paragraph(safe(line), qa_a_style))
                continue
            if current_section == "Failure Signals":
                clean = line.lstrip("-•").strip()
                story.append(Paragraph("•  " + safe(clean), bullet_style))
                continue
            if current_section == "Key Takeaways":
                story.append(Paragraph("→  " + safe(line), takeaway_style))
                continue
            if current_section == "Practical Reference":
                if line.startswith("T-Code:"):
                    story.append(Paragraph(safe(line), tcode_style))
                    continue
                if (line.startswith("$") or line.startswith(">") or
                    line.startswith("SELECT") or line.startswith("db.") or
                    line.startswith("aws ") or line.startswith("docker ") or
                    line.startswith("kubectl") or line.startswith("pip ") or
                    line.startswith("python") or line.startswith("def ") or
                    line.startswith("import ") or line.startswith("curl ") or
                    line.startswith("nmap ") or line.startswith("git ") or
                    (len(line) < 120 and any(c in line for c in ["()", "[]", "{}", "=>", "->"]))):
                    story.append(Paragraph(safe(line), code_style))
                    continue
            if current_section in SUBHEAD_SECTIONS:
                if len(line) < 80 and not line.endswith('.') and not line.endswith(','):
                    story.append(Paragraph(safe(line), subhead_style))
                    continue
            story.append(Paragraph(safe(line), body_style))
            if current_section in ("Real-Time Examples", "Practical Reference"):
                story.append(light_hr())

        story.append(Spacer(1, 0.3*inch))
        story.append(HRFlowable(width="100%", thickness=1,
            color=colors.HexColor('#1F5C99')))
        story.append(Paragraph(
            "iMentora Training Platform  |  " + safe_domain + "  |  Class Notes",
            footer_style
        ))
        doc.build(story)

        # FIX: verify file actually written
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            logger.info(f"[PDF] Generated successfully | size={os.path.getsize(output_path)} bytes")
            return True
        else:
            logger.error(f"[PDF] File missing or empty after build: {output_path}")
            return False

    except Exception as e:
        logger.error(f"[PDF ERROR] {e}", exc_info=True)
        return False


def _generate_transcript_pdf(transcript_text: str, domain: str, output_path: str) -> bool:
    """Generates transcript as PDF. Returns True on success."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        doc = SimpleDocTemplate(
            output_path, pagesize=letter,
            rightMargin=0.75*inch, leftMargin=0.75*inch,
            topMargin=0.75*inch, bottomMargin=0.75*inch
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('TTitle', parent=styles['Title'],
            fontSize=18, textColor=colors.HexColor('#1F5C99'),
            spaceAfter=4, alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('TSubtitle', parent=styles['Normal'],
            fontSize=10, textColor=colors.HexColor('#555555'),
            spaceAfter=16, alignment=TA_CENTER)
        body_style = ParagraphStyle('TBody', parent=styles['Normal'],
            fontSize=10, leading=16, spaceAfter=6,
            alignment=TA_JUSTIFY, textColor=colors.HexColor('#222222'))
        footer_style = ParagraphStyle('TFooter', parent=styles['Normal'],
            fontSize=8, textColor=colors.HexColor('#888888'),
            alignment=TA_CENTER, spaceBefore=5)

        def safe(text):
            return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        story = []
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(safe(domain) + " — Full Transcript", title_style))
        story.append(Paragraph("iMentora Training Session", subtitle_style))
        story.append(HRFlowable(width="100%", thickness=2,
            color=colors.HexColor('#1F5C99'), spaceAfter=12))
        story.append(Spacer(1, 0.1*inch))
        words = transcript_text.split()
        for i in range(0, len(words), 80):
            chunk = " ".join(words[i:i + 80])
            if chunk.strip():
                story.append(Paragraph(safe(chunk), body_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1F5C99')))
        story.append(Paragraph(
            "iMentora Training Platform  |  " + safe(domain) + "  |  Full Transcript",
            footer_style
        ))
        doc.build(story)

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            logger.info(f"[TRANSCRIPT PDF] Generated | size={os.path.getsize(output_path)} bytes")
            return True
        else:
            logger.error("[TRANSCRIPT PDF] File missing or empty")
            return False
    except Exception as e:
        logger.error(f"[TRANSCRIPT PDF ERROR] {e}", exc_info=True)
        return False


def save_docx(content: str, path: str, image_path: str = None, title: str = ""):
    from docx.shared import RGBColor
    doc = Document()
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            run = doc.add_heading(line[4:].strip(), level=3).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith("## "):
            run = doc.add_heading(line[3:].strip(), level=2).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith("# "):
            run = doc.add_heading(line[2:].strip(), level=1).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.lower().startswith("example"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)
    if image_path and os.path.exists(image_path):
        doc.add_page_break()
        doc.add_heading("Mind Map", level=2)
        doc.add_picture(image_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)

async def transcribe_chunk(chunk_file: str, offset: float):
    try:
        with open(chunk_file, "rb") as f:
            result = groq_client.audio.transcriptions.create(
                model="whisper-large-v3",
                file=f,
                response_format="verbose_json"
            )
        segments = []
        if hasattr(result, "segments") and result.segments:
            for seg in result.segments:
                segments.append({
                    "start": offset + float(seg["start"]),
                    "end": offset + float(seg["end"]),
                    "text": seg["text"].strip()
                })
        else:
            text = (result.text or "").strip()
            if not text:
                text = "[No speech detected]"
            segments.append({"start": offset, "end": offset + 5, "text": text})
        return segments
    except Exception as e:
        logger.error(f"[GROQ ERROR] {e}")
        return [{"start": offset, "end": offset + 5, "text": "[Transcription failed]"}]


# ══════════════════════════════════════════════════════════
# DOMAIN DETECTION
# ══════════════════════════════════════════════════════════

def detect_domain_and_type(transcript: str) -> tuple:
    """
    Single LLM call — detects both domain and user_type.
    Works for ANY domain without hardcoded lists.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{
                "role": "user",
                "content": f"""
Analyze this transcript and return JSON with two fields:

1. "domain" — the MAIN TECHNICAL SUBJECT being taught (not the admin chat
   at the start). Look for the primary technical topic, such as:
   "Cybersecurity CVE CVSS CWE", "SAP HANA", "SAP MM", "Python",
   "Java", "AWS", "Digital Marketing", "SAP SD", "Selenium", "HR", etc.
   IMPORTANT: Ignore introductory admin talk like recording issues,
   system setup, VirtualBox installation — focus on what is actually
   being TAUGHT as the subject matter of the class.

2. "user_type" — either "dev" or "non_dev" using this rule:
   - "dev" = people who WRITE CODE or BUILD SYSTEMS
     Examples: Python, Java, JavaScript, React, Node, AWS, DevOps,
     Docker, Kubernetes, QA Automation, Selenium, Data Science, ML,
     Cybersecurity (coding exploits), Networking, Linux, SQL development
   - "non_dev" = people who USE SYSTEMS or MANAGE BUSINESS PROCESSES
     Examples: ANY SAP module (HANA, MM, SD, FICO, Basis, BW, HR),
     Digital Marketing, CRM, Salesforce, ServiceNow, HR processes,
     Finance, Procurement, Business Analysis, IT Support, Helpdesk,
     Cybersecurity CVE/CVSS/CWE analysis (not coding), vulnerability
     management, security compliance, manual testing

CRITICAL RULES:
- SAP anything → ALWAYS non_dev, no exceptions
- CVE, CVSS, CWE, vulnerability analysis → non_dev
- Writing actual code / building APIs / automation scripts → dev
- QA manual testing → non_dev
- QA automation / Selenium / coding tests → dev
- If class is about vulnerability identification/scoring/reporting → non_dev
- If class mentions CVE IDs, CVSS scores, NVD website → domain is Cybersecurity

Reply with ONLY valid JSON. No explanation. No markdown.
Example: {{"domain": "Cybersecurity CVE CVSS CWE", "user_type": "non_dev"}}

TRANSCRIPT (first 4000 chars — read fully before deciding):
{transcript[:4000]}
"""
            }],
            temperature=0.0,
            max_tokens=50
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r"```json|```", "", raw).strip()
        data = json.loads(raw)
        domain    = data.get("domain", "General IT")
        user_type = data.get("user_type", "dev")
        logger.info(f"🎯 Domain: {domain} | 👤 Type: {user_type}")
        return domain, user_type
    except Exception as e:
        logger.error(f"Domain detection failed: {e}")
        return "General IT", "dev"


def _strip_placeholders(text: str) -> str:
    """Remove any leaked placeholder text from LLM output."""
    text = re.sub(r'\*?Suggested next steps:.*?(\n|$)', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Diagram Placeholder.*?(\n|$)',       '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[Insert.*?\]',                      '', text, flags=re.IGNORECASE)
    text = re.sub(r'Mind Map.*?(\n|$)',                  '', text, flags=re.IGNORECASE)
    text = re.sub(r'```dot.*?```',                       '', text, flags=re.DOTALL | re.IGNORECASE)
    return text.strip()


# ══════════════════════════════════════════════════════════
# WEB ENRICHMENT
# ══════════════════════════════════════════════════════════


def _translate_transcript_to_english(transcript: str) -> str:
    """
    If transcript contains significant non-English content
    (Telugu, Hindi, Tamil etc), translate it to English
    before summarizing. This gives the summary LLM clean
    English to work with instead of mixed languages.
    Technical terms like CVE, CVSS, SAP, Python stay as-is.
    """
    try:
        # Quick check — count non-ASCII characters
        non_ascii = sum(1 for c in transcript if ord(c) > 127)
        ratio = non_ascii / max(len(transcript), 1)

        if ratio < 0.1:
            # Less than 10% non-ASCII — mostly English already
            logger.info("Transcript is mostly English — skipping translation")
            return transcript

        logger.info(f"Non-ASCII ratio: {ratio:.2f} — translating to English")

        # Translate in chunks of 4000 chars to stay within token limits
        chunks = [transcript[i:i+4000]
                  for i in range(0, min(len(transcript), 32000), 4000)]
        translated_chunks = []

        for chunk in chunks:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": f"""Translate this training transcript to English.

RULES:
- Keep all technical terms in English as-is
  (CVE, CVSS, CWE, SAP, T-codes, Python, SQL, AWS etc)
- Keep the conversational classroom style — do not clean it up
- Translate only the non-English words to English
- Preserve every example and every concept the trainer mentioned
- Do not summarize — translate the full content

TRANSCRIPT:
{chunk}

Return ONLY the translated text. Nothing else."""
                }],
                temperature=0.0,
                max_tokens=4000
            )
            translated_chunks.append(
                response.choices[0].message.content.strip()
            )

        translated = " ".join(translated_chunks)
        logger.info(
            f"✅ Translation complete | "
            f"original={len(transcript)} | translated={len(translated)}"
        )
        return translated

    except Exception as e:
        logger.warning(f"Translation failed — using original: {e}")
        return transcript

def _extract_topics(transcript: str) -> list:
    """
    Extract key topics from transcript using word frequency.
    No hardcoded keywords — works for any domain.
    """
    STOPWORDS = {
        'the', 'and', 'for', 'are', 'this', 'that', 'with',
        'from', 'have', 'been', 'will', 'they', 'your', 'into',
        'more', 'also', 'each', 'some', 'when', 'which', 'their',
        'there', 'these', 'those', 'would', 'could', 'should',
        'where', 'while', 'about', 'after', 'before', 'other',
        'using', 'being', 'having', 'through', 'following',
        'example', 'based', 'during', 'within', 'without',
        'process', 'system', 'result', 'status', 'value',
        'level', 'error', 'check', 'output', 'input', 'state',
    }
    words = re.findall(r'\b[a-z]{5,}\b', transcript.lower())
    counts = Counter(w for w in words if w not in STOPWORDS)
    topics = [w for w, _ in counts.most_common(10)]
    logger.info(f"📚 Extracted topics: {topics[:5]}")
    return topics


def _fetch_web_context(domain: str, topics: list) -> str:
    """
    Search web for real-world content about the domain and topics.
    All queries built dynamically — nothing hardcoded.
    """
    try:
        from duckduckgo_search import DDGS

        queries = [
            f"{domain} common issues support engineer",
            f"{domain} troubleshooting scenarios real world",
            f"{domain} error messages root cause",
        ]
        for topic in topics[:3]:
            queries.append(f"{domain} {topic} problem resolution")

        all_content = []
        with DDGS() as ddgs:
            for query in queries[:5]:
                try:
                    results = ddgs.text(query, max_results=3)
                    for r in results:
                        snippet = r.get("body", "")
                        if snippet and len(snippet) > 100:
                            all_content.append(snippet)
                except Exception:
                    continue

        if not all_content:
            return ""

        combined = "\n\n".join(all_content[:10])
        logger.info(f"🌐 Web context fetched | domain={domain} | chars={len(combined)}")
        return combined

    except Exception as e:
        logger.warning(f"Web fetch failed: {e}")
        return ""


# ══════════════════════════════════════════════════════════
# DEV SUMMARY — single definition, class notebook style
# ══════════════════════════════════════════════════════════

def _generate_dev_summary(transcript: str, domain: str, context: str) -> str:
    """
    Expert-level class notebook for senior engineers.
    No hardcoded domains or team names.
    Includes: session notes, real-time examples, failure signals,
    errors, mistakes, takeaways, interview questions, glossary.
    """
    prompt = f"""
IMPORTANT: The transcript below may be in Telugu, Hindi, Tamil,
or any other Indian language mixed with English.
Read and understand the transcript fully in whatever language it is.
Write the ENTIRE OUTPUT in English only — not Telugu, not Hindi.
Use the transcript content to identify exactly what topics
the trainer covered today. Do not ignore non-English content.

You are a SENIOR ENGINEER with 5+ years of hands-on production
experience in {domain}.

You are writing a CLASS NOTEBOOK that combines what the trainer
taught today with your own deep field experience. Write like a
senior engineer who attended the class and enriched the notes
with real production insights, war stories, counter-intuitive
lessons, and the kind of knowledge that only comes from years
of working in real enterprise environments.

MINIMUM OUTPUT: 25000 characters. Do not stop early. Write until you reach 25000 characters minimum.

DOMAIN: {domain}
TRANSCRIPT (today's class content):
{transcript[:20000]}
{context}

══════════════════════════════════════════════════════
WRITE THESE SECTIONS IN ORDER
Use each heading exactly as shown on its own line.
Write plain paragraphs under each heading.
No bullet points except in Failure Signals section.
No markdown. No bold. No #. No **.
══════════════════════════════════════════════════════

Overview

Write 2-3 lines only.
What today's class is about.
What domain and topic was covered.
Why this topic matters in real production work.

Today's Session Notes

This is the MOST IMPORTANT section. Write it long and detailed.
Write exactly like a student taking running notes in class.
Include every concept the trainer explained.
Include every example the trainer gave — especially:
  Any CVE IDs the trainer showed or typed on screen
  Any CVSS scores the trainer calculated or demonstrated
  Any websites the trainer opened and showed live
  Any vulnerability names like Log4j, Heartbleed, etc.
  Any score ranges the trainer wrote or explained
  Any tools the trainer demonstrated
Include every technical term the trainer used and explain
each term in plain English immediately after it appears.
If the trainer explained a process step by step write every
step with what happens at each step and why it matters.
If the trainer demonstrated something on screen — describe
exactly what was on the screen, what was typed, what appeared.
If the trainer gave a real-world scenario write it fully
with all details including exact numbers and identifiers.
This section must be at least 8000 characters on its own.
Write in flowing paragraphs — like notes not a report.

Core Concepts

Write one paragraph per key concept from today's class.
For each concept write:
What the concept is in simple terms first.
Then the technical explanation with correct terminology.
Then a real example from {domain} showing how it works.
Then what happens when this concept is misunderstood.
Minimum 4 concepts. Each paragraph minimum 6 sentences.
Every concept must have a real example — not a generic one.

For each concept you MUST include:
The exact system behavior at a technical level — not just
what it does but HOW it does it internally.
The exact log message, error text, or system output that
appears when this concept fails in production.
The cascading effect on other components when this fails —
what breaks next and in what order.
The metric or signal that appears BEFORE the failure becomes
visible to the user — the early warning indicator.
Why the obvious fix sometimes makes things worse — the
counter-intuitive behavior that only experience reveals.
The specific parameter, configuration key, transaction code,
command, or system table that controls this behavior.

IMPORTANT FOR CORE CONCEPTS:
Use ONLY what the trainer taught — pull exact details from
the transcript. If the trainer showed CVE-2021-44228 with
score 10 — use that exact example. If the trainer explained
CVSS metrics (Attack Vector, Confidentiality, Integrity,
Availability) — include those exact metrics. Do not replace
transcript content with generic textbook explanations.



Real-Time Examples

Write one real example for every concept covered today.
Each example must be a real situation in {domain} production.
Format each example like this:
In a real {domain} environment — [describe the exact situation,
what the engineer sees, what they do, what happens, and why
this example illustrates the concept].
Every example must be different from the others.
Minimum 5 examples. Each example minimum 4 sentences.

DOMAIN-SPECIFIC EXAMPLE RULES:
If {domain} involves Python, Java, JavaScript, Node, React or
any programming language — every example must show the actual
error message from terminal or logs, the exact line of code
that caused it, and the corrected code.
If {domain} involves AWS, Docker, Kubernetes, Linux, DevOps —
every example must show the exact CLI command used, the exact
output or error returned, and the correct follow-up command.
If {domain} involves SQL, MongoDB, databases — every example
must show the exact query that failed or was slow, the explain
plan output or error, and the corrected query.
If {domain} involves Cybersecurity, CVE, CVSS, penetration
testing — every example must show the actual CVE ID from the
transcript (do NOT invent a CVE like CVE-2023-12345), the CVSS
score the trainer mentioned, the exact tool or URL shown, and
what was demonstrated on screen.
If {domain} involves SAP — every example must show the exact
transaction code used, the exact field values entered, and
what the system showed or what error appeared.
If {domain} involves Selenium, QA Automation — every example
must show the actual test script, the exact assertion that
failed, and the corrected test code.

CRITICAL: Never invent CVE IDs, CWE IDs, or tool commands.
Only use exact identifiers and URLs the trainer showed in class.

Practical Reference

This section is a hands-on reference card for {domain}.
Write based on what was taught today.

DOMAIN-SPECIFIC RULES — follow exactly:

If {domain} involves Python, Java, JavaScript, Node, React,
or any programming language — write the actual code snippets
from today's class. For each snippet write what it does,
the exact code, what output it produces, and what error it
throws if misused. Minimum 5 code examples. Each must be
runnable and real. Show exact error messages from terminal.

If {domain} involves AWS, Docker, Kubernetes, Linux, DevOps,
shell scripting, or any CLI tool — write the exact commands
covered today. For each command write what it does, the exact
syntax, a real example with values filled in, and expected
output or error. Minimum 5 commands. Copy-paste ready.

If {domain} involves SQL, MongoDB, Redis, or any database —
write the exact queries covered today. For each query write
what it does, the exact syntax, a real example with table or
collection names, and the output. Minimum 5 queries. Each
must be executable as-is.

If {domain} involves SAP — any module (MM, SD, FICO, HANA,
Basis, Security, BW, HR, CRM) — write the exact transaction
codes covered today. For each T-code write exactly this:
T-Code: [code] — [what it does] — Navigation: [menu path]
Use: [when to use it] — Fields: [key fields to fill]
Minimum 8 transaction codes from today's class.

If {domain} involves Cybersecurity, CVE, CVSS, CWE,
penetration testing, ethical hacking, or vulnerability
analysis — write the exact commands, CVE lookup steps,
and CVSS calculation examples from today. For each write
the exact command or URL, what it returns, and how to
interpret the output. Minimum 5 actionable steps.

If {domain} involves Selenium, QA Automation, or test
frameworks — write the exact test scripts and locator
strategies from today. For each write the test scenario,
the exact script, what it tests, and what failure output
looks like. Minimum 5 complete runnable examples.

If {domain} involves Digital Marketing, HR, CRM, Salesforce,
ServiceNow, or any business tool — write the exact process
steps, form fields, and navigation paths from today. For each
write the step name, exact navigation, fields to fill, and
expected outcome. Minimum 5 actionable process steps.

STRICT RULES FOR THIS SECTION:
- Every single item must come DIRECTLY from the transcript.
- If the trainer showed a T-code, write that exact T-code.
- If the trainer ran a command, write that exact command.
- If the trainer wrote code, write that exact code.
- If the trainer showed a CVE ID, write that exact CVE ID.
- Do NOT invent examples that are not in the transcript.
- Do NOT write generic examples from your own knowledge.
- Do NOT guess the domain and add typical examples.

EXAMPLE OF WHAT NOT TO DO:
If the class was about CVE/CVSS but the trainer never used
SAP T-codes in this class — do NOT write SU01, SE80, SM37
or any SAP T-codes. Those are invented. That is wrong.

If the class was about CVE/CVSS and the trainer showed
CVE-2021-44228 with score 10 and went to nvd.nist.gov —
write ONLY that CVE ID, that score, and that URL.
Nothing else. Do not add T-codes or commands not shown.

EXAMPLE OF WHAT TO DO:
CVE-2021-44228 — Log4j vulnerability — CVSS score: 10.0 (Critical)
URL shown in class: https://nvd.nist.gov/vuln/detail/CVE-2021-44228
What the trainer showed: the description, the score breakdown,
the attack vector, confidentiality/integrity/availability ratings.

- If the transcript does not contain enough practical content
  for this section, write only what exists and state clearly:
  "The trainer focused on concepts in this session.
   Practical commands will be covered in the next class."

Write this section as plain paragraphs.
No bullet points. No markdown.


Failure Signals

Write this section as bullet points — this is the only exception.
List every warning sign, alert, or signal that indicates
something is going wrong in {domain}.
For each signal write:
- [Signal name] — [what it looks like] — [what it means]
Minimum 10 signals. Keep each line short and sharp.
Group them by category if there are many.

Common Errors

Write one paragraph per error type from {domain}.
For each error write:
The exact error name or message.
What causes this error in production.
A real example of when this error appears.
What the correct fix is and why.
What the wrong fix is and what it causes.
Minimum 5 errors. Write like a reference guide a student
can come back to when they see this error in real work.

Common Mistakes Students Make

Write one paragraph per common mistake.
These are mistakes that students and junior engineers
make when they first start working in {domain}.
For each mistake write:
What the student does wrong.
Why it seems correct at first.
What actually happens as a result.
What the correct approach is.
Minimum 4 mistakes.

Key Takeaways

Write the most important things to remember from today's class.
Write like the trainer is summarising the last 5 minutes of class.
Each takeaway must be a practical point — not theory.
Example style: Always check X before Y because...
Example style: Never do Z in production because...
Example style: When you see X it always means Y not Z...
Minimum 8 takeaways. Write as short punchy paragraphs.

Interview Questions

Write 8 interview questions based exactly on today's content.
These must be real questions asked in {domain} interviews.
Mix these types:
2 scenario questions — you see X happening, what do you do
2 error diagnosis questions — this error appears, what is the cause
2 process questions — what happens before or after this step
2 concept questions — explain X in context of Y situation
For each question write the question then the answer in
one clear paragraph below it.
Questions must test real understanding — not definition recall.

Terminology Glossary

Write every technical term that appeared in today's class.
For each term write:
The term name.
What it means in plain English.
How it is used in real {domain} work.
One example sentence showing it in context.
Minimum 10 terms. Write in paragraph style — not a list.

Escalation Guide

Based on {domain}, derive the real teams that handle
different types of problems. Do not hardcode team names —
derive them from the domain.
For each team write what problems go to them, what
information to collect before escalating, what not to
escalate to them, and what happens when you go to the
wrong team.
Minimum 3 teams.

Best Practices

Write one paragraph per best practice from today's class.
Each best practice must come from something covered today —
not generic advice.
For each practice write:
What the practice is.
A real situation where skipping it caused a problem.
The correct approach with reasoning.
Minimum 4 best practices.

══════════════════════════════════════════════════════
OUTPUT RULES
══════════════════════════════════════════════════════

- Write each section heading on its own line exactly as shown
- Plain paragraphs under every heading except Failure Signals
- Failure Signals section uses bullet points only
- No markdown — no #, **, backtick characters anywhere
- No Suggested next steps
- No Diagram Placeholder
- No Prerequisites section
- No Conclusion section
- Minimum 25000 characters total — mandatory
- Today's Session Notes must be minimum 8000 characters
- End after Best Practices, nothing else
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": (
                    "The transcript may be in Telugu, Hindi, Tamil, or mixed languages. "
                    "Read it fully and understand it. Write all output in English only. "
                    f"You are a brilliant student and note-taker attending a {domain} "
                    "training class. You write the most detailed class notes possible. "
                    "You include every example, every technical term, every concept "
                    "the trainer covered. "
                    "You write Today's Session Notes like running class notes — "
                    "long, detailed, with every example and technical term explained. "
                    "You write Common Errors like a reference guide with real examples. "
                    "You write Interview Questions mixing scenario, error diagnosis, "
                    "process, and concept question types. "
                    "You write Terminology Glossary covering every term from class. "
                    f"You derive all escalation teams from {domain} — never hardcode. "
                    "You always write minimum 25000 characters. You expand every section with deep detail, multiple real examples, thorough explanations. You never stop early. "
                    "You never write bullet points except in Failure Signals section. "
                    "You never write markdown, bold, or # or ** anywhere. "
                    "You always complete every section before stopping. "
                    "You write Today's Session Notes with extreme detail — minimum 8000 characters in that section alone. "
                    "You write every concept with 8+ sentences. You write 8+ interview questions. "
                    "You write 12+ failure signals. You write 6+ common errors. You write 12+ takeaways. "
                    "You write 12+ glossary terms. You never truncate any section."
                )
            },
            {"role": "user", "content": prompt}
        ],
        temperature=0.5,
        max_tokens=16000
    )
    return response.choices[0].message.content.strip()


# ══════════════════════════════════════════════════════════
# NON-DEV SUMMARY — single definition, class notebook style
# ══════════════════════════════════════════════════════════

def _generate_non_dev_summary(transcript: str, domain: str, context: str) -> str:
    """
    Expert-level class notebook for senior QA, BA, Support professionals.
    No hardcoded domains or team names.
    Includes: session notes, real-time examples, failure signals,
    errors, mistakes, takeaways, interview questions, glossary.
    """
    prompt = f"""
IMPORTANT: The transcript below may be in Telugu, Hindi, Tamil,
or any other Indian language mixed with English.
Read and understand the transcript fully in whatever language it is.
Write the ENTIRE OUTPUT in English only — not Telugu, not Hindi.
Use the transcript content to identify exactly what topics
the trainer covered today. Do not ignore non-English content.

You are a SENIOR ENTERPRISE SUPPORT PROFESSIONAL with 5+ years
of hands-on experience in {domain}.

You are writing a CLASS NOTEBOOK that combines what the trainer
taught today with your own deep field experience. Write like a
senior support engineer who attended the class and enriched the
notes with real case studies, actual escalation paths, production
war stories, and the kind of knowledge that separates a junior
from a senior in real enterprise environments.

MINIMUM OUTPUT: 25000 characters. Do not stop early. Write until you reach 25000 characters minimum.

DOMAIN: {domain}
TRANSCRIPT (today's class content):
{transcript[:20000]}
{context}

══════════════════════════════════════════════════════
WRITE THESE SECTIONS IN ORDER
Use each heading exactly as shown on its own line.
Write plain paragraphs under each heading.
No bullet points except in Failure Signals section.
No markdown. No bold. No #. No **.
══════════════════════════════════════════════════════

Overview

Write 2-3 lines only.
What today's class is about.
What domain and topic was covered.
Why this topic matters in real enterprise support work.

Today's Session Notes

This is the MOST IMPORTANT section. Write it long and detailed.
Write exactly like a student taking running notes in class.
Include every concept the trainer explained.
Include every business process the trainer described.
Include every real-world example the trainer gave.
Include every technical term used and explain each term
in plain English immediately after it appears.
If the trainer described a business process write the
complete process flow — every step, what happens at each
step, what goes wrong at each step, and why it matters.
If the trainer gave a support scenario write it fully
with all the signals, the diagnosis, and the resolution.
This section must be at least 8000 characters on its own.
Write in flowing paragraphs — like notes not a report.

Core Concepts

Write one paragraph per key concept from today's class.
For each concept write:
What the concept means in simple business terms first.
Then the precise technical or business explanation.
Then a real example from {domain} showing how it works
in an actual enterprise environment.
Then what happens when this concept is misunderstood.
Minimum 4 concepts. Each paragraph minimum 6 sentences.
Every concept must have a real example tied to {domain}.

For each concept you MUST include:
The exact system behavior at a technical level — not just
what it does but HOW it does it internally.
The exact log message, error text, or system output that
appears when this concept fails in production.
The cascading effect on other components when this fails —
what breaks next and in what order.
The metric or signal that appears BEFORE the failure becomes
visible to the user — the early warning indicator.
Why the obvious fix sometimes makes things worse — the
counter-intuitive behavior that only experience reveals.
The specific parameter, configuration key, transaction code,
command, or system table that controls this behavior.

IMPORTANT FOR CORE CONCEPTS:
Use ONLY what the trainer taught — pull exact details from
the transcript. If the trainer showed CVE-2021-44228 with
score 10 — use that exact example. If the trainer explained
CVSS metrics (Attack Vector, Confidentiality, Integrity,
Availability) — include those exact metrics. Do not replace
transcript content with generic textbook explanations.



Real-Time Examples

Write one real example for every concept covered today.
Each example must be a real enterprise situation in {domain}.
Format each example like this:
In a real {domain} environment — [describe the exact business
situation, what the support engineer or analyst sees,
what they do, what happens, and why this example
illustrates the concept clearly].
Every example must be specific — not generic.
Minimum 5 examples. Each example minimum 4 sentences.

DOMAIN-SPECIFIC EXAMPLE RULES:
If {domain} involves Python, Java, JavaScript, Node, React or
any programming language — every example must show the actual
error message from terminal or logs, the exact line of code
that caused it, and the corrected code.
If {domain} involves AWS, Docker, Kubernetes, Linux, DevOps —
every example must show the exact CLI command used, the exact
output or error returned, and the correct follow-up command.
If {domain} involves SQL, MongoDB, databases — every example
must show the exact query that failed or was slow, the explain
plan output or error, and the corrected query.
If {domain} involves Cybersecurity, CVE, CVSS, penetration
testing — every example must show the actual CVE ID from the
transcript (do NOT invent a CVE like CVE-2023-12345), the CVSS
score the trainer mentioned, the exact tool or URL shown, and
what was demonstrated on screen.
If {domain} involves SAP — every example must show the exact
transaction code used, the exact field values entered, and
what the system showed or what error appeared.
If {domain} involves Selenium, QA Automation — every example
must show the actual test script, the exact assertion that
failed, and the corrected test code.

CRITICAL: Never invent CVE IDs, CWE IDs, or tool commands.
Only use exact identifiers and URLs the trainer showed in class.

Practical Reference

This section is a hands-on reference card for {domain}.
Write based on what was taught today.

DOMAIN-SPECIFIC RULES — follow exactly:

If {domain} involves Python, Java, JavaScript, Node, React,
or any programming language — write the actual code snippets
from today's class. For each snippet write what it does,
the exact code, what output it produces, and what error it
throws if misused. Minimum 5 code examples. Each must be
runnable and real. Show exact error messages from terminal.

If {domain} involves AWS, Docker, Kubernetes, Linux, DevOps,
shell scripting, or any CLI tool — write the exact commands
covered today. For each command write what it does, the exact
syntax, a real example with values filled in, and expected
output or error. Minimum 5 commands. Copy-paste ready.

If {domain} involves SQL, MongoDB, Redis, or any database —
write the exact queries covered today. For each query write
what it does, the exact syntax, a real example with table or
collection names, and the output. Minimum 5 queries. Each
must be executable as-is.

If {domain} involves SAP — any module (MM, SD, FICO, HANA,
Basis, Security, BW, HR, CRM) — write the exact transaction
codes covered today. For each T-code write exactly this:
T-Code: [code] — [what it does] — Navigation: [menu path]
Use: [when to use it] — Fields: [key fields to fill]
Minimum 8 transaction codes from today's class.

If {domain} involves Cybersecurity, CVE, CVSS, CWE,
penetration testing, ethical hacking, or vulnerability
analysis — write the exact commands, CVE lookup steps,
and CVSS calculation examples from today. For each write
the exact command or URL, what it returns, and how to
interpret the output. Minimum 5 actionable steps.

If {domain} involves Selenium, QA Automation, or test
frameworks — write the exact test scripts and locator
strategies from today. For each write the test scenario,
the exact script, what it tests, and what failure output
looks like. Minimum 5 complete runnable examples.

If {domain} involves Digital Marketing, HR, CRM, Salesforce,
ServiceNow, or any business tool — write the exact process
steps, form fields, and navigation paths from today. For each
write the step name, exact navigation, fields to fill, and
expected outcome. Minimum 5 actionable process steps.

STRICT RULES FOR THIS SECTION:
- Every single item must come DIRECTLY from the transcript.
- If the trainer showed a T-code, write that exact T-code.
- If the trainer ran a command, write that exact command.
- If the trainer wrote code, write that exact code.
- If the trainer showed a CVE ID, write that exact CVE ID.
- Do NOT invent examples that are not in the transcript.
- Do NOT write generic examples from your own knowledge.
- Do NOT guess the domain and add typical examples.

EXAMPLE OF WHAT NOT TO DO:
If the class was about CVE/CVSS but the trainer never used
SAP T-codes in this class — do NOT write SU01, SE80, SM37
or any SAP T-codes. Those are invented. That is wrong.

If the class was about CVE/CVSS and the trainer showed
CVE-2021-44228 with score 10 and went to nvd.nist.gov —
write ONLY that CVE ID, that score, and that URL.
Nothing else. Do not add T-codes or commands not shown.

EXAMPLE OF WHAT TO DO:
CVE-2021-44228 — Log4j vulnerability — CVSS score: 10.0 (Critical)
URL shown in class: https://nvd.nist.gov/vuln/detail/CVE-2021-44228
What the trainer showed: the description, the score breakdown,
the attack vector, confidentiality/integrity/availability ratings.

- If the transcript does not contain enough practical content
  for this section, write only what exists and state clearly:
  "The trainer focused on concepts in this session.
   Practical commands will be covered in the next class."

Write this section as plain paragraphs.
No bullet points. No markdown.


Failure Signals

Write this section as bullet points — this is the only exception.
List every warning sign, system alert, business indicator,
or user complaint that signals something is going wrong
in {domain}.
For each signal write:
- [Signal] — [what it looks like in the system or ticket] — [what it means]
Minimum 10 signals. Keep each line short and sharp.
Group by category if there are many types.

Common Errors

Write one paragraph per error or issue type from {domain}.
For each error write:
What the error or issue is called.
What causes it in a real enterprise environment.
A real example of when it appears — what the user reports
or what the system shows.
What the correct resolution is and why.
What the wrong approach is and what it causes.
Minimum 5 errors. Write like a reference guide the student
can use when they see this issue in real support work.

Common Mistakes Students Make

Write one paragraph per common mistake.
These are mistakes that junior support engineers, QA analysts,
or BAs make when they first start working in {domain}.
For each mistake write:
What the student does wrong.
Why it seems like the right thing to do at first.
What actually happens as a result in the business.
What the correct approach is.
Minimum 4 mistakes.

Key Takeaways

Write the most important things to remember from today's class.
Write like the trainer is summarising the last 5 minutes.
Each takeaway must be practical — not theory.
Example style: Always check X before escalating because...
Example style: When you see X in the system it means Y not Z...
Example style: Never do X during month-end close because...
Minimum 8 takeaways. Write as short punchy paragraphs.

Interview Questions

Write 8 interview questions based exactly on today's content.
These must be real questions asked in {domain} interviews.
Mix these types:
2 scenario questions — a user reports X and Y is also showing, what do you check first
2 error questions — this error appears in the system, what is the root cause
2 process questions — what happens before or after this step in the workflow
2 concept questions — explain X in the context of a real business situation
For each question write the question then the answer in
one clear paragraph below it.
Questions must test real understanding — not definition recall.

Terminology Glossary

Write every technical and business term from today's class.
For each term write:
The term.
What it means in plain English for someone new to {domain}.
How it is used in real enterprise {domain} work.
One example sentence showing it in context.
Minimum 10 terms. Write in paragraph style — not a list.

Escalation Guide

Based on {domain}, derive the real teams that handle
different types of problems in an enterprise environment.
Do not hardcode any team names — derive from {domain}.
For each team write what problems go to them, what
information to collect before escalating, what NOT to
escalate to them, and consequence of wrong escalation.
Minimum 3 teams.

Best Practices

Write one paragraph per best practice from today's class.
Each best practice must connect to something the trainer
covered — not generic advice.
For each write:
What the practice is.
A real situation where skipping it caused a business problem.
The correct approach with reasoning.
Minimum 4 best practices.

══════════════════════════════════════════════════════
OUTPUT RULES
══════════════════════════════════════════════════════

- Write each section heading on its own line exactly as shown
- Plain paragraphs under every heading except Failure Signals
- Failure Signals section uses bullet points only
- No commands or file paths anywhere
- No markdown — no #, **, backtick characters anywhere
- No Suggested next steps
- No Diagram Placeholder
- No Prerequisites section
- No Conclusion section
- Minimum 25000 characters total — mandatory
- Today's Session Notes must be minimum 8000 characters
- End after Best Practices, nothing else
"""
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": (
                    "The transcript may be in Telugu, Hindi, Tamil, or mixed languages. "
                    "Read it fully and understand it. Write all output in English only. "
                    f"You are a brilliant student and note-taker attending a {domain} "
                    "training class for enterprise support professionals. "
                    "You write the most detailed class notes possible. "
                    "You include every business process, every example, every technical "
                    "term, every support scenario the trainer covered. "
                    "You write Today's Session Notes like running class notes — "
                    "long, detailed, with every example and every term explained. "
                    "You write Common Errors like a real support reference guide. "
                    "You write Interview Questions mixing scenario, error, process, "
                    "and concept question types. "
                    "You write Terminology Glossary covering every term from class. "
                    f"You derive all escalation teams from {domain} — never hardcode. "
                    "You always write minimum 25000 characters. You expand every section with deep detail, multiple real examples, thorough explanations. You never stop early. "
                    "You never write bullet points except in Failure Signals section. "
                    "You never write markdown, bold, or # or ** anywhere. "
                    "You always complete every section before stopping. "
                    "You write Today's Session Notes with extreme detail — minimum 8000 characters in that section alone. "
                    "You write every concept with 8+ sentences. You write 8+ interview questions. "
                    "You write 12+ failure signals. You write 6+ common errors. You write 12+ takeaways. "
                    "You write 12+ glossary terms. You never truncate any section."
                )
            },
            {"role": "user", "content": prompt}
        ],
        temperature=0.5,
        max_tokens=16000
    )
    return response.choices[0].message.content.strip()


def _deep_enrich_from_transcript(transcript: str, domain: str) -> str:
    """
    Makes a second GPT-4o-mini call to extract deeper content from the
    transcript — real examples, commands, exact terms, analogies the
    trainer used — and returns it as additional context for the main
    summary LLM. This replaces web search enrichment and works fully
    offline. Output is ~5000 chars of structured extra content.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a meticulous note extraction assistant. "
                        "Your job is to pull out every concrete detail from "
                        "a training transcript — exact commands, exact tool names, "
                        "exact examples, exact analogies, exact numbers, exact URLs, "
                        "exact error messages, exact steps the trainer demonstrated. "
                        "Write in plain English. No markdown. No bullet headers. "
                        "Just dense, detailed paragraphs of extracted content. "
                        "Minimum 5000 characters."
                    )
                },
                {
                    "role": "user",
                    "content": f"""Read this training transcript carefully and extract ALL of the following:

DOMAIN: {domain}
TRANSCRIPT:
{transcript[:25000]}

Extract and write in detail:

1. EXACT COMMANDS AND TOOLS: Every command, tool name, software, URL, or
   website the trainer mentioned or demonstrated. Write the exact syntax
   if given. For example if trainer said "nmap -sS" write "nmap -sS".
   If trainer showed a website write the exact URL.

2. EXACT EXAMPLES FROM CLASS: Every specific example the trainer gave.
   Include all numbers, IDs, names, scores mentioned. For example if
   trainer said "CVE-2021-44228 has score 10" write that exactly.
   If trainer showed a government website as example, name it exactly.

3. STEP BY STEP PROCESSES: Every process or workflow the trainer
   explained step by step. Write all steps in order with what happens
   at each step.

4. ANALOGIES AND EXPLANATIONS: Every analogy or plain-English explanation
   the trainer used to explain a concept. For example "Google index is
   like a ledger" — write the full analogy as the trainer explained it.

5. STUDENT INTERACTIONS: What questions students asked and how the
   trainer answered. What practical exercises students did.

6. WARNINGS AND EMPHASIS: Everything the trainer said students must
   remember, must not do, or emphasized strongly.

7. REAL-WORLD SCENARIOS: Every real production scenario, real company
   example, or real incident the trainer described.

8. TECHNICAL DETAILS: Every technical detail — port numbers, protocol
   names, configuration parameters, file names, error messages.

Write everything you find. Be exhaustive. Do not summarize — extract.
Minimum 5000 characters of extracted content."""
                }
            ],
            temperature=0.2,
            max_tokens=6000
        )
        result = response.choices[0].message.content.strip()
        logger.info(f"[DEEP ENRICH] Extracted {len(result)} chars from transcript")
        return result
    except Exception as e:
        logger.warning(f"[DEEP ENRICH] Failed: {e}")
        return ""


# ══════════════════════════════════════════════════════════
# MAIN SUMMARIZATION ENTRY POINT
# ══════════════════════════════════════════════════════════

def summarize_segment(transcript: str, context: str = "") -> dict:
    """
    1. Detects domain AND user_type  (GPT-4o-mini — no hardcoded lists)
    2. Extracts key topics from transcript (word frequency — no hardcoded keywords)
    3. Fetches web content using domain + topics (dynamic queries — no hardcoded domains)
    4. Generates class notebook summary using transcript + web content combined
    5. Strips all placeholder text
    6. Returns dict with summary, user_type, domain

    dev      → db["Developer"]
    non_dev  → db["Non-Developer"]
    """
    # Step 1 — translate transcript to English FIRST if needed
    english_transcript = _translate_transcript_to_english(transcript)

    # Step 1.5 — detect domain and user_type from ENGLISH transcript
    domain, user_type = detect_domain_and_type(english_transcript)
    logger.info(f"👤 User type: {user_type} | Domain: {domain}")

    # Step 2 — extract topics from translated transcript
    topics = _extract_topics(english_transcript)

    # Step 3 — deep transcript enrichment (no internet needed)
    deep_context = _deep_enrich_from_transcript(english_transcript, domain)
    logger.info(f"✅ Deep enrichment generated | chars={len(deep_context)}")

    # Step 4 — build enriched context
    enriched_context = context
    if deep_context:
        enriched_context = (
            context +
            "\n\nADDITIONAL CONTEXT EXTRACTED FROM TRANSCRIPT "
            "(use this to write richer, more detailed notes):\n"
            + deep_context
        )

    # Step 5 — generate summary using english_transcript + deep context
    try:
        if user_type == "dev":
            summary = _generate_dev_summary(english_transcript, domain, enriched_context)
        else:
            summary = _generate_non_dev_summary(english_transcript, domain, enriched_context)
    except Exception as e:
        logger.error(f"[ERROR] Summary generation failed: {e}")
        summary = "Summary generation failed."

    # Step 6 — strip placeholders
    summary = _strip_placeholders(summary)

    logger.info(
        f"✅ Summary ready | domain={domain} | "
        f"user_type={user_type} | chars={len(summary)}"
    )

    return {
        "summary":   summary,
        "user_type": user_type,
        "domain":    domain,
    }


# ══════════════════════════════════════════════════════════
# TRAINER PERFORMANCE EVALUATION
# ══════════════════════════════════════════════════════════

def analyze_trainer_performance(transcript: str) -> dict:
    if not transcript.strip():
        return {
            "technical_content": 0,
            "explanation_clarity": 0,
            "friendliness": 0,
            "communication": 0,
            "overall_feedback": "No speech detected for evaluation."
        }

    prompt = f"""
You are an expert communication and training evaluator.
Evaluate the trainer's communication quality, tone, and content in the transcript below.

TRANSCRIPT:
\"\"\"{ transcript}\"\"\"

Evaluate across the following dimensions (each scored 0-100%):
1. Technical Content — accuracy, depth, and domain clarity.
2. Explanation Clarity — how logically and simply ideas are explained.
3. Friendliness — warmth, politeness, and positive tone.
4. Communication — evaluate using Indian English standards:
   - Focus on fluency, confidence, and comfort with Indian accent.
   - Accept light Indianisms such as "basically", "ok na", "ya", etc.
   - Penalize only unclear speech or excessive filler use.
   - Do not reduce marks for accent style.
Output a short JSON report with numeric scores and 1-2 lines of feedback.
Format exactly as:
{{
  "technical_content": <number>,
  "explanation_clarity": <number>,
  "friendliness": <number>,
  "communication": <number>,
  "overall_feedback": "<short summary>"
}}
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a behavioral analytics and technical communication evaluator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )
        raw = response.choices[0].message.content.strip()
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        else:
            logger.warning("Unable to parse trainer evaluation JSON.")
            return {"error": "Failed to parse evaluation output."}
    except Exception as e:
        logger.error(f"[ERROR] Trainer performance evaluation failed: {e}")
        return {"error": str(e)}


# ══════════════════════════════════════════════════════════
# VIDEO PROCESS  — ALL 5 FIXES APPLIED HERE
# ══════════════════════════════════════════════════════════

async def process_video(
    video_path: str,
    meeting_id: str,
    user_id:    str,
    batch_id:   str = None,
    course:     str = None
):
    with TemporaryDirectory() as workdir:
        compressed = os.path.join(workdir, "compressed.mp4")
        audio_path = os.path.join(workdir, "audio.wav")

        # 1 — Compression + Noise Cancellation
        subprocess.run([
            "ffmpeg", "-y", "-i", video_path,
            "-af", "afftdn=nf=-25",
            "-c:v", "libx264", "-crf", "35", "-preset", "ultrafast",
            "-c:a", "aac", "-b:a", "64k", compressed
        ], check=True)

        # 2 — Extract Audio
        subprocess.run([
            "ffmpeg", "-y", "-i", compressed,
            "-ar", "16000", "-ac", "1", "-vn", audio_path
        ], check=True)

        # 3 — Transcription
        audio = AudioSegment.from_wav(audio_path)
        chunk_length = 5 * 60 * 1000
        audio_chunks = []
        for i in range(0, len(audio), chunk_length):
            chunk_file = os.path.join(workdir, f"chunk_{i//chunk_length}.wav")
            chunk = audio[i:i + chunk_length]
            chunk.export(chunk_file, format="wav")
            offset = audio[:i].duration_seconds
            audio_chunks.append((chunk_file, offset))

        tasks = [transcribe_chunk(path, offset) for path, offset in audio_chunks]
        chunk_results = await asyncio.gather(*tasks)
        all_segments = [seg for result in chunk_results for seg in result]
        full_transcript = " ".join(seg['text'] for seg in all_segments)

        # 3.1 — Trainer Performance Evaluation
        trainer_scores = analyze_trainer_performance(full_transcript)

        # 4 — Subtitles
        srt_paths = {}
        for lang in ["en", "hi", "te"]:
            translated = []
            for seg in all_segments:
                text = seg.get("text", "")
                if lang != "en":
                    try:
                        text = GoogleTranslator(source='en', target=lang).translate(text)
                    except Exception:
                        text = "[Translation failed]"
                translated.append({
                    "start": seg.get("start", 0),
                    "end": seg.get("end", 0),
                    "text": text
                })
            srt_path = os.path.join(workdir, f"subs_{lang}.srt")
            create_srt_from_segments(translated, srt_path)
            srt_paths[lang] = srt_path

        print("Subtitle check:")
        print(f"  Exists: {os.path.exists(srt_paths['en'])}")
        if os.path.exists(srt_paths['en']):
            print(f"  Size: {os.path.getsize(srt_paths['en'])} bytes")

        # 5 — Captioned Video
        captioned = os.path.join(workdir, "captioned.mp4")
        srt_path_fixed = os.path.abspath(srt_paths["en"]).replace("\\", "/")
        if not os.path.exists(srt_path_fixed):
            raise FileNotFoundError(f"Subtitle file missing: {srt_path_fixed}")

        cmd = [
            "ffmpeg", "-y",
            "-i", compressed,
            "-vf", f"subtitles={srt_path_fixed}:force_style='FontName=Arial\\,FontSize=18'",
            "-c:v", "libx264", "-c:a", "aac",
            captioned
        ]
        subprocess.run(cmd, check=True)

        # 6 — Summary Generation (transcript + web enrichment)
        summary_result   = summarize_segment(full_transcript)
        summary_text_raw = summary_result["summary"]
        user_type        = summary_result["user_type"]
        domain           = summary_result["domain"]

        collection = dev_collection if user_type == "dev" else non_dev_collection
        logger.info(f"Storing in: {'Developer' if user_type == 'dev' else 'Non-Developer'} collection")

        # 7 — Mind Map Extraction
        image_path = os.path.join(workdir, "mindmap.png")
        image_url = None

        dot_match = re.search(r"```(?:dot)?\s*(.*?)```", summary_text_raw, re.DOTALL)
        if not dot_match:
            dot_match = re.search(r"(digraph\s+[A-Za-z0-9_]*\s*\{.*?\})", summary_text_raw, re.DOTALL)

        if dot_match:
            dot_code = dot_match.group(1).strip()
            try:
                generate_graph(dot_code, image_path[:-4])
                image_url = upload_to_s3(
                    "summary-image",
                    image_path,
                    f"{meeting_id}_{user_id}_mindmap.png"
                )
            except Exception as e:
                logger.error(f"[GRAPH ERROR] {e}")
                image_url = None

        # 7.1 — Remove DOT code before cleaning
        summary_text_raw_no_dot = re.sub(r"```(?:dot)?\s*.*?```", "", summary_text_raw, flags=re.DOTALL)
        summary_text_raw_no_dot = re.sub(r"digraph\s+[A-Za-z0-9_]*\s*\{.*?\}", "", summary_text_raw_no_dot, flags=re.DOTALL)

        # 8 — Clean markdown
        summary_text = clean_markdown(summary_text_raw_no_dot)

        # FIX 1: Generate Summary PDF and capture return value
        summary_pdf = os.path.join(workdir, "summary.pdf")
        logger.info(f"[PDF] Generating at: {summary_pdf}")
        pdf_ok = generate_summary_pdf(summary_text, domain, summary_pdf)
        logger.info(f"[PDF] Result: {'SUCCESS' if pdf_ok else 'FAILED'}")

        # FIX 2: Generate Transcript as PDF instead of DOCX
        transcript_pdf = os.path.join(workdir, "transcript.pdf")
        logger.info(f"[TRANSCRIPT PDF] Generating at: {transcript_pdf}")
        transcript_pdf_ok = _generate_transcript_pdf(full_transcript, domain, transcript_pdf)
        logger.info(f"[TRANSCRIPT PDF] Result: {'SUCCESS' if transcript_pdf_ok else 'FAILED'}")

        # 9 — Upload to S3
        video_url = upload_to_s3("videos", captioned,
                                 f"{meeting_id}_{user_id}_captioned.mp4")

        # FIX 3: Upload transcript as PDF, fall back to DOCX only if PDF failed
        transcript_url = ""
        if transcript_pdf_ok and os.path.exists(transcript_pdf):
            transcript_url = upload_to_s3("transcripts", transcript_pdf,
                                          f"{meeting_id}_{user_id}_transcript.pdf")
            logger.info(f"[TRANSCRIPT PDF] Uploaded -> {transcript_url}")
        else:
            transcript_docx = os.path.join(workdir, "transcript.docx")
            save_docx(full_transcript, transcript_docx, title="Full Transcript")
            transcript_url = upload_to_s3("transcripts", transcript_docx,
                                          f"{meeting_id}_{user_id}_transcript.docx")
            logger.warning(f"[TRANSCRIPT] Fell back to DOCX -> {transcript_url}")

        # FIX 4: Only upload summary PDF if file actually exists and has content
        summary_url = ""
        if pdf_ok and os.path.exists(summary_pdf) and os.path.getsize(summary_pdf) > 0:
            summary_url = upload_to_s3("summary", summary_pdf,
                                       f"{meeting_id}_{user_id}_summary.pdf")
            logger.info(f"[PDF] Summary uploaded -> {summary_url}")
        else:
            logger.error(f"[PDF] Skipping upload — file missing or empty: {summary_pdf}")

        summary_pdf_url = summary_url

        subtitle_urls = {}
        for lang, path in srt_paths.items():
            subtitle_urls[lang] = upload_to_s3(
                "videos", path, f"{meeting_id}_{user_id}_subs_{lang}.srt"
            )

        # 10 — Save to MongoDB
        collection.insert_one({
            "meeting_id":         meeting_id,
            "user_id":            user_id,
            "batch_id":           batch_id,
            "course":             course,
            "filename":           os.path.basename(video_path),
            "domain":             domain,
            "user_type":          user_type,
            "video_url":          video_url,
            "transcript_url":     transcript_url,
            "summary_url":        summary_url,
            "summary_pdf_url":    summary_pdf_url,
            "image_url":          image_url,
            "subtitles":          subtitle_urls,
            "transcript_text":    full_transcript,
            "summary":            summary_text,
            "trainer_evaluation": trainer_scores,
            "timestamp":          datetime.now()
        })

        logger.info(
            f"[DONE] meeting={meeting_id} | "
            f"summary_pdf={summary_url} | transcript={transcript_url}"
        )

        # FIX 5: Return domain in response so UI can display it
        return {
            "status":            "success",
            "domain":            domain,
            "video_url":         video_url,
            "transcript_url":    transcript_url,
            "summary_url":       summary_url,
            "summary_pdf_url":   summary_pdf_url,
            "summary_image_url": image_url,
            "subtitle_urls":     subtitle_urls,
        }


# ══════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════

@app.get("/recordings")
async def get_recordings(user_id: str, collection: str = None):
    def sanitize(obj):
        """Recursively convert any datetime or non-serializable value to JSON-safe type."""
        if isinstance(obj, datetime):
            return obj.isoformat()
        if isinstance(obj, dict):
            return {k: sanitize(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [sanitize(i) for i in obj]
        try:
            import json
            json.dumps(obj)
            return obj
        except (TypeError, ValueError):
            return str(obj)

    try:
        query = {"user_id": user_id}
        results = []
        if collection == "Developer":
            for doc in dev_collection.find(query, {"_id": 0}):
                doc["_collection"] = "Developer"
                results.append(sanitize(doc))
        elif collection == "Non-Developer":
            for doc in non_dev_collection.find(query, {"_id": 0}):
                doc["_collection"] = "Non-Developer"
                results.append(sanitize(doc))
        else:
            for col, name in [(dev_collection, "Developer"), (non_dev_collection, "Non-Developer")]:
                for doc in col.find(query, {"_id": 0}):
                    doc["_collection"] = name
                    results.append(sanitize(doc))
        return JSONResponse(content={"recordings": results})
    except Exception as e:
        logger.error(f"[RECORDINGS] Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload/")
async def upload_single(
    file:       UploadFile = File(...),
    meeting_id: str        = Form(...),
    user_id:    str        = Form(...),
    batch_id:   str        = Form(None),
    course:     str        = Form(None)
):
    try:
        existing = (
            dev_collection.find_one({"meeting_id": meeting_id, "user_id": user_id, "filename": file.filename})
            or
            non_dev_collection.find_one({"meeting_id": meeting_id, "user_id": user_id, "filename": file.filename})
        )
        if existing:
            ts = existing.get("timestamp")
            if isinstance(ts, datetime):
                ts = ts.timestamp()
            return {
                "status":            "already_processed",
                "file":              file.filename,
                "domain":            existing.get("domain", ""),
                "video_url":         existing.get("video_url"),
                "transcript_url":    existing.get("transcript_url"),
                "summary_url":       existing.get("summary_url") or existing.get("summary_pdf_url", ""),
                "summary_pdf_url":   existing.get("summary_url") or existing.get("summary_pdf_url", ""),
                "summary_image_url": existing.get("image_url"),
                "subtitle_urls":     existing.get("subtitles"),
                "message":           "Already processed."
            }

        with TemporaryDirectory() as tmp:
            temp_path = os.path.join(tmp, file.filename)
            with open(temp_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
            result = await process_video(
                temp_path, meeting_id, user_id,
                batch_id=batch_id, course=course
            )
            result["file"] = file.filename
            logger.info(f"[UPLOAD] summary_url={result.get('summary_url')} transcript_url={result.get('transcript_url')}")
            return JSONResponse(content=result)

    except Exception as e:
        logger.exception("Upload failed")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/")
def home():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/health")
def health():
    return {"status": "healthy"}

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8010, reload=True)